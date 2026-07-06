"""Утилиты для обработки мультимодальных вложений в чате."""

from __future__ import annotations

import base64
from io import BytesIO

from docx import Document
from fastapi import UploadFile
from openai import AsyncOpenAI
from pypdf import PdfReader

_MAX_TEXT_LENGTH = 30_000
_MAX_PDF_PAGES = 50
_openai_client: AsyncOpenAI | None = None


def _get_openai_client() -> AsyncOpenAI:
    """Лениво создает OpenAI-клиент для расшифровки голосовых сообщений."""

    global _openai_client
    if _openai_client is None:
        _openai_client = AsyncOpenAI()
    return _openai_client


async def media_to_part(media: UploadFile) -> dict:
    """Преобразует загруженный файл в content-part для OpenAI Chat Completions."""

    mime = media.content_type or ""
    data = await media.read()

    if mime.startswith("image/"):
        b64 = base64.b64encode(data).decode("ascii")
        return {
            "type": "image_url",
            "image_url": {"url": f"data:{mime};base64,{b64}"},
        }

    if mime.startswith("audio/") or mime == "application/ogg":
        transcript = await whisper_transcribe(data, media.filename or "audio.ogg")
        return {
            "type": "text",
            "text": f"[пользователь сказал голосом]:\n{transcript}",
        }

    if mime == "application/pdf":
        text = extract_pdf_text(data)[:_MAX_TEXT_LENGTH]
        return {"type": "text", "text": f"[документ PDF]:\n{text}"}

    if mime.endswith("wordprocessingml.document"):
        text = extract_docx_text(data)[:_MAX_TEXT_LENGTH]
        return {"type": "text", "text": f"[документ DOCX]:\n{text}"}

    raise ValueError(f"Unsupported media type: {mime}")


async def whisper_transcribe(audio_bytes: bytes, filename: str) -> str:
    """Отправляет аудио в Whisper и возвращает готовую расшифровку."""

    file_like = BytesIO(audio_bytes)
    file_like.name = filename
    result = await _get_openai_client().audio.transcriptions.create(
        model="whisper-1",
        file=file_like,
    )
    return result.text


def extract_pdf_text(data: bytes) -> str:
    """Извлекает текст из PDF и помечает вероятно сканированные документы."""

    reader = PdfReader(BytesIO(data))
    fragments: list[str] = []
    sparse_pages = 0

    for page in reader.pages[:_MAX_PDF_PAGES]:
        text = (page.extract_text() or "").strip()
        fragments.append(text)
        if len(text) < 100:
            sparse_pages += 1

    merged = "\n\n".join(fragment for fragment in fragments if fragment)
    if len(reader.pages) >= 5 and sparse_pages >= 5:
        note = "Документ похож на сканированный PDF, текст извлечен частично."
        return f"{note}\n\n{merged}".strip()
    return merged


def extract_docx_text(data: bytes) -> str:
    """Извлекает текст из DOCX, включая содержимое таблиц."""

    document = Document(BytesIO(data))
    chunks: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                chunks.append(" | ".join(cells))

    return "\n".join(chunks)
